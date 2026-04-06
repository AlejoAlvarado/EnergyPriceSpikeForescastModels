from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "Data" / "CSVs" / "aeso_merged_2020_2025.csv"
OUTPUT_ROOT = ROOT / "output"
FIGURES_DIR = OUTPUT_ROOT / "figures"
REPORT_DIR = OUTPUT_ROOT / "report"
DOCX_DIR = OUTPUT_ROOT / "docx"
PDF_DIR = OUTPUT_ROOT / "pdf"

for directory in (FIGURES_DIR, REPORT_DIR, DOCX_DIR, PDF_DIR):
    directory.mkdir(parents=True, exist_ok=True)


TITLE = "Predicting Alberta Electricity Price Spikes: Evidence From Supply, Demand, and Renewable Conditions"
AUTHORS = "Aravindh Palaniguru, Alejandro Jose Alvarado Barrera, and Jorge Gutierrez Barajas"
COURSE = "DATA 607"
REPORT_DATE = "April 5, 2026"

PIPELINE_SUMMARY = [
    ["MLP", "59", "48,887", "33,696", "9,648", "5,543", "0.91"],
    ["LSTM", "32", "48,839", "33,672", "9,624", "5,543", "0.94"],
    ["CNN", "92", "48,839", "33,672", "9,624", "5,543", "0.77"],
]

MODEL_SUMMARY = [
    ["Naive", "0.384", "0.384", "0.384", "n/a"],
    ["MLP", "0.363", "0.275", "0.535", "0.934"],
    ["LSTM", "0.394", "0.311", "0.535", "0.945"],
    ["CNN", "0.416", "0.340", "0.535", "0.941"],
]

REFERENCES = [
    "Alberta Electric System Operator. (2025a). Hourly metered volumes and pool price and AIL data 2001 to July 2025 [Data set]. https://www.aeso.ca/market/market-and-system-reporting/data-requests/hourly-generation-metered-volumes-and-pool-price-and-ail-data-2001-to-july-2025/",
    "Alberta Electric System Operator. (2025b). Historical generation data (CSD) [Data set]. https://www.aeso.ca/market/market-and-system-reporting/data-requests/historical-generation-data/",
    "Lago, J., De Ridder, F., & De Schutter, B. (2018). Forecasting spot electricity prices: Deep learning approaches and empirical comparison of traditional algorithms. Applied Energy, 221, 386-405.",
    "Manfre Jaimes, D., Zamudio Lopez, M., Zareipour, H., & Quashie, M. (2023). A hybrid model for multi-day-ahead electricity price forecasting considering price spikes. Forecasting, 5(3), 499-521.",
    "Saha, C. (2025). Developing a statistical risk assessment and grid prediction tool for power system reliability (Master's capstone project). University of Calgary. https://ucalgary.scholaris.ca/server/api/core/bitstreams/2660b5fd-7457-4720-b2c1-f4127a579d62/content",
    "Zamudio Lopez, M., et al. (2024). Forecasting the occurrence of electricity price spikes: A statistical-economic investigation study. Forecasting, 6(1), 7.",
]


@dataclass
class FigureSpec:
    key: str
    title: str
    filename: str
    note: str

    @property
    def path(self) -> Path:
        return FIGURES_DIR / self.filename


FIGURE_SPECS = [
    FigureSpec(
        key="price_dynamics",
        title="Distribution and temporal structure of Alberta pool prices.",
        filename="figure_1_price_dynamics.png",
        note=(
            "Left panel shows the heavy right tail in hourly pool prices. Right panel shows mean hourly prices "
            "by month and hour of day, highlighting seasonal and intraday structure."
        ),
    ),
    FigureSpec(
        key="spike_conditions",
        title="Mean system conditions for hours with and without a future spike at t+2.",
        filename="figure_2_spike_conditions.png",
        note=(
            "Future spike hours exhibit higher current prices and net load, but lower wind output, lower "
            "renewables share, and lower reserve margin."
        ),
    ),
    FigureSpec(
        key="correlations",
        title="Correlation matrix for selected market and system-stress variables.",
        filename="figure_3_correlation_matrix.png",
        note=(
            "The exploratory matrix summarizes pairwise linear association and is used only to describe structure, "
            "not to make causal claims."
        ),
    ),
    FigureSpec(
        key="model_results",
        title="Test performance comparison reported in Data607_EnergySpike_Presentation_v2.pptx.",
        filename="figure_4_model_results.png",
        note=(
            "Panel A compares the rounded test F1 values reported in the presentation deck. Panel B shows that "
            "precision improves from MLP to LSTM to CNN while recall remains at approximately 0.535, so most of "
            "the F1 gain comes from reducing false positives rather than capturing additional spike hours."
        ),
    ),
]


def load_dataframe() -> pd.DataFrame:
    df = pd.read_csv(DATA_PATH, parse_dates=["datetime"])
    df["month"] = df["datetime"].dt.month
    df["hour"] = df["datetime"].dt.hour
    df["year"] = df["datetime"].dt.year
    return df


def save_price_dynamics(df: pd.DataFrame, spec: FigureSpec) -> None:
    sns.set_theme(style="whitegrid")
    fig, axes = plt.subplots(1, 2, figsize=(12, 4.5))

    sns.histplot(df["ACTUAL_POOL_PRICE"], bins=80, ax=axes[0], color="#b31b1b")
    axes[0].set_title("Hourly pool price distribution")
    axes[0].set_xlabel("Pool price (CAD/MWh)")
    axes[0].set_ylabel("Count")
    axes[0].text(
        0.97,
        0.95,
        "Median = 45.34\np95 = 431.00\nMax = 999.99",
        transform=axes[0].transAxes,
        ha="right",
        va="top",
        fontsize=10,
        bbox={"boxstyle": "round,pad=0.3", "facecolor": "white", "edgecolor": "#999999"},
    )

    pivot = df.pivot_table(index="month", columns="hour", values="ACTUAL_POOL_PRICE", aggfunc="mean")
    sns.heatmap(pivot, ax=axes[1], cmap="mako", cbar_kws={"label": "Mean price"})
    axes[1].set_title("Mean hourly price by month and hour")
    axes[1].set_xlabel("Hour of day")
    axes[1].set_ylabel("Month")

    fig.tight_layout()
    fig.savefig(spec.path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def save_spike_conditions(df: pd.DataFrame, spec: FigureSpec) -> None:
    grouped = df.groupby("spike_lead_2")[
        ["ACTUAL_POOL_PRICE", "net_load", "wind_total", "renewables_share", "reserve_margin"]
    ].mean()
    labels = {0: "No future spike", 1: "Future spike at t+2"}
    cols = [
        ("ACTUAL_POOL_PRICE", "Current pool price\n(CAD/MWh)"),
        ("net_load", "Net load\n(MW)"),
        ("wind_total", "Wind output\n(MW)"),
        ("renewables_share", "Renewables share"),
        ("reserve_margin", "Reserve margin"),
    ]

    sns.set_theme(style="whitegrid")
    fig, axes = plt.subplots(1, len(cols), figsize=(13.5, 3.6))
    palette = ["#7f8c8d", "#b31b1b"]

    for ax, (column, title) in zip(axes, cols):
        values = grouped[column]
        ax.bar([labels[0], labels[1]], [values.loc[0], values.loc[1]], color=palette, width=0.65)
        ax.set_title(title, fontsize=10)
        ax.tick_params(axis="x", rotation=35)
        for idx, value in enumerate([values.loc[0], values.loc[1]]):
            text = f"{value:.3f}" if value < 1 else f"{value:,.0f}"
            ax.text(idx, value, text, ha="center", va="bottom", fontsize=8)

    fig.tight_layout()
    fig.savefig(spec.path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def save_correlations(df: pd.DataFrame, spec: FigureSpec) -> None:
    cols = [
        "ACTUAL_POOL_PRICE",
        "ACTUAL_AIL",
        "net_load",
        "wind_total",
        "solar_total",
        "gas_total",
        "renewables_share",
        "reserve_margin",
        "net_load_3h_change",
        "spike_lead_2",
    ]
    corr = df[cols].corr(numeric_only=True)

    sns.set_theme(style="white")
    fig, ax = plt.subplots(figsize=(8.5, 6.5))
    sns.heatmap(
        corr,
        annot=True,
        fmt=".2f",
        cmap="coolwarm",
        center=0,
        linewidths=0.5,
        cbar_kws={"label": "Pearson r"},
        ax=ax,
    )
    ax.set_title("Selected variable correlations")
    fig.tight_layout()
    fig.savefig(spec.path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def save_model_results(spec: FigureSpec) -> None:
    sns.set_theme(style="whitegrid")
    fig, axes = plt.subplots(1, 2, figsize=(11.5, 4.5))

    f1_models = ["Naive", "MLP", "LSTM", "CNN"]
    f1_values = [0.384, 0.363, 0.394, 0.416]
    axes[0].bar(f1_models, f1_values, color=["#95a5a6", "#5d6d7e", "#2874a6", "#b31b1b"])
    axes[0].set_ylim(0, 0.6)
    axes[0].set_title("Test F1-score")
    axes[0].set_ylabel("F1")
    for idx, value in enumerate(f1_values):
        axes[0].text(idx, value + 0.01, f"{value:.3f}", ha="center", fontsize=9)

    models = ["MLP", "LSTM", "CNN"]
    precision = [0.275, 0.311, 0.340]
    recall = [0.535, 0.535, 0.535]
    x = range(len(models))
    width = 0.36
    axes[1].bar([i - width / 2 for i in x], precision, width=width, color="#5d6d7e", label="Precision")
    axes[1].bar([i + width / 2 for i in x], recall, width=width, color="#b31b1b", label="Recall")
    axes[1].set_xticks(list(x))
    axes[1].set_xticklabels(models)
    axes[1].set_ylim(0, 0.7)
    axes[1].set_title("Test precision and recall")
    axes[1].set_ylabel("Score")
    axes[1].legend(frameon=False, loc="upper left")
    for idx, value in enumerate(precision):
        axes[1].text(idx - width / 2, value + 0.015, f"{value:.3f}", ha="center", fontsize=9)
    for idx, value in enumerate(recall):
        axes[1].text(idx + width / 2, value + 0.015, f"{value:.3f}", ha="center", fontsize=9)

    fig.tight_layout()
    fig.savefig(spec.path, dpi=300, bbox_inches="tight")
    plt.close(fig)


def generate_figures(df: pd.DataFrame) -> None:
    save_price_dynamics(df, FIGURE_SPECS[0])
    save_spike_conditions(df, FIGURE_SPECS[1])
    save_correlations(df, FIGURE_SPECS[2])
    save_model_results(FIGURE_SPECS[3])


def add_page_number(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_docx_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)


def add_centered_paragraph(document: Document, text: str, bold: bool = False, size: int = 12) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def add_figure_block(document: Document, figure: FigureSpec) -> None:
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(f"Figure {FIGURE_SPECS.index(figure) + 1}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    title_run = cap.add_run(figure.title)
    title_run.italic = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)

    document.add_picture(str(figure.path), width=Inches(6.2))
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = note.add_run(f"Note. {figure.note}")
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_table_block(document: Document, number: int, title: str, headers: list[str], rows: list[list[str]]) -> None:
    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = cap.add_run(f"Table {number}. ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    title_run = cap.add_run(title)
    title_run.italic = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(12)

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def add_reference_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.5)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def build_docx(report_text: dict[str, list[str]]) -> Path:
    output_path = DOCX_DIR / "data607_project_report_pptx_metrics.docx"
    doc = Document()
    configure_docx_styles(doc)

    for _ in range(7):
        doc.add_paragraph()
    add_centered_paragraph(doc, TITLE, bold=True, size=14)
    doc.add_paragraph()
    add_centered_paragraph(doc, AUTHORS)
    add_centered_paragraph(doc, COURSE)
    add_centered_paragraph(doc, REPORT_DATE)

    doc.add_page_break()
    add_heading(doc, "Abstract")
    add_body_paragraph(doc, report_text["abstract"][0])

    add_heading(doc, "Background and Introduction")
    for paragraph in report_text["introduction"]:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "Data Source and Preparation")
    for paragraph in report_text["data"]:
        add_body_paragraph(doc, paragraph)
    add_table_block(
        doc,
        1,
        "Model pipeline overview taken from the source notebooks.",
        ["Model", "Features", "Rows", "Train", "Validation", "Test", "Threshold"],
        PIPELINE_SUMMARY,
    )

    add_heading(doc, "Preliminary Analyses")
    for paragraph in report_text["eda"]:
        add_body_paragraph(doc, paragraph)
    add_figure_block(doc, FIGURE_SPECS[0])
    add_figure_block(doc, FIGURE_SPECS[1])
    add_figure_block(doc, FIGURE_SPECS[2])

    add_heading(doc, "Problem Statement and Working Hypotheses")
    for paragraph in report_text["problem"]:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "Formal Analyses")
    for paragraph in report_text["methods"]:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "Results and Interpretation")
    for paragraph in report_text["results"]:
        add_body_paragraph(doc, paragraph)
    add_table_block(
        doc,
        2,
        "Rounded test metrics reported in Data607_EnergySpike_Presentation_v2.pptx.",
        ["Model", "Test F1", "Precision", "Recall", "ROC-AUC"],
        MODEL_SUMMARY,
    )
    add_figure_block(doc, FIGURE_SPECS[3])

    add_heading(doc, "Conclusions")
    for paragraph in report_text["conclusion"]:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "Task Division")
    for paragraph in report_text["task_division"]:
        add_body_paragraph(doc, paragraph)

    add_heading(doc, "References")
    for reference in REFERENCES:
        add_reference_paragraph(doc, reference)

    doc.save(output_path)
    return output_path


def pdf_styles():
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TitleAPA",
            parent=base["Title"],
            fontName="Times-Bold",
            fontSize=16,
            leading=20,
            alignment=TA_CENTER,
            spaceAfter=18,
        ),
        "center": ParagraphStyle(
            "CenterAPA",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=18,
            alignment=TA_CENTER,
            spaceAfter=6,
        ),
        "heading1": ParagraphStyle(
            "Heading1APA",
            parent=base["Normal"],
            fontName="Times-Bold",
            fontSize=12,
            leading=18,
            alignment=TA_CENTER,
            spaceBefore=12,
            spaceAfter=6,
        ),
        "body": ParagraphStyle(
            "BodyAPA",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=12,
            leading=18,
            alignment=TA_JUSTIFY,
            spaceAfter=6,
        ),
        "caption": ParagraphStyle(
            "CaptionAPA",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=11,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=6,
        ),
        "note": ParagraphStyle(
            "NoteAPA",
            parent=base["Normal"],
            fontName="Times-Italic",
            fontSize=10,
            leading=12,
            alignment=TA_LEFT,
            spaceAfter=8,
        ),
        "reference": ParagraphStyle(
            "ReferenceAPA",
            parent=base["Normal"],
            fontName="Times-Roman",
            fontSize=11,
            leading=14,
            leftIndent=0.5 * inch,
            firstLineIndent=-0.5 * inch,
            spaceAfter=4,
        ),
    }


def build_pdf_table(headers: list[str], rows: list[list[str]], widths: list[float]) -> Table:
    base_styles = getSampleStyleSheet()
    header_style = ParagraphStyle(
        "TableHeader",
        parent=base_styles["Normal"],
        fontName="Times-Bold",
        fontSize=8.5,
        leading=10,
        alignment=TA_LEFT,
    )
    body_style = ParagraphStyle(
        "TableBody",
        parent=base_styles["Normal"],
        fontName="Times-Roman",
        fontSize=8,
        leading=10,
        alignment=TA_LEFT,
    )

    data = [[Paragraph(header, header_style) for header in headers]]
    for row in rows:
        data.append([Paragraph(value, body_style) for value in row])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
            ]
        )
    )
    return table


def pdf_page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Times-Roman", 10)
    canvas.drawRightString(LETTER[0] - doc.rightMargin, LETTER[1] - 0.6 * inch, str(canvas.getPageNumber()))
    canvas.restoreState()


def build_pdf(report_text: dict[str, list[str]]) -> Path:
    output_path = PDF_DIR / "data607_project_report_pptx_metrics.pdf"
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    story = [
        Spacer(1, 2.1 * inch),
        Paragraph(TITLE, styles["title"]),
        Paragraph(AUTHORS, styles["center"]),
        Paragraph(COURSE, styles["center"]),
        Paragraph(REPORT_DATE, styles["center"]),
        PageBreak(),
        Paragraph("Abstract", styles["heading1"]),
        Paragraph(report_text["abstract"][0], styles["body"]),
        Paragraph("Background and Introduction", styles["heading1"]),
    ]
    for paragraph in report_text["introduction"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Data Source and Preparation", styles["heading1"]))
    for paragraph in report_text["data"]:
        story.append(Paragraph(paragraph, styles["body"]))
    story.extend(
        [
            Paragraph("<b>Table 1.</b> <i>Model pipeline overview taken from the source notebooks.</i>", styles["caption"]),
            build_pdf_table(
                ["Model", "Features", "Rows", "Train", "Validation", "Test", "Threshold"],
                PIPELINE_SUMMARY,
                [0.75 * inch, 0.7 * inch, 0.8 * inch, 0.85 * inch, 0.95 * inch, 0.8 * inch, 0.75 * inch],
            ),
            Spacer(1, 0.18 * inch),
            Paragraph("Preliminary Analyses", styles["heading1"]),
        ]
    )
    for paragraph in report_text["eda"]:
        story.append(Paragraph(paragraph, styles["body"]))
    for figure in FIGURE_SPECS[:3]:
        story.append(
            KeepTogether(
                [
                    Paragraph(
                        f"<b>Figure {FIGURE_SPECS.index(figure) + 1}.</b> <i>{figure.title}</i>",
                        styles["caption"],
                    ),
                    RLImage(str(figure.path), width=6.1 * inch, height=3.35 * inch),
                    Paragraph(f"Note. {figure.note}", styles["note"]),
                ]
            )
        )

    story.append(Paragraph("Problem Statement and Working Hypotheses", styles["heading1"]))
    for paragraph in report_text["problem"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Formal Analyses", styles["heading1"]))
    for paragraph in report_text["methods"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Results and Interpretation", styles["heading1"]))
    for paragraph in report_text["results"]:
        story.append(Paragraph(paragraph, styles["body"]))
    story.extend(
        [
            Paragraph("<b>Table 2.</b> <i>Rounded test metrics reported in Data607_EnergySpike_Presentation_v2.pptx.</i>", styles["caption"]),
            build_pdf_table(
                ["Model", "Test F1", "Precision", "Recall", "ROC-AUC"],
                MODEL_SUMMARY,
                [1.0 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch, 1.0 * inch],
            ),
            Spacer(1, 0.18 * inch),
        ]
    )
    story.append(
        KeepTogether(
            [
                Paragraph(f"<b>Figure 4.</b> <i>{FIGURE_SPECS[3].title}</i>", styles["caption"]),
                RLImage(str(FIGURE_SPECS[3].path), width=6.1 * inch, height=3.25 * inch),
                Paragraph(f"Note. {FIGURE_SPECS[3].note}", styles["note"]),
            ]
        )
    )

    story.append(Paragraph("Conclusions", styles["heading1"]))
    for paragraph in report_text["conclusion"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("Task Division", styles["heading1"]))
    for paragraph in report_text["task_division"]:
        story.append(Paragraph(paragraph, styles["body"]))

    story.append(Paragraph("References", styles["heading1"]))
    for reference in REFERENCES:
        story.append(Paragraph(reference, styles["reference"]))

    doc.build(story, onFirstPage=pdf_page_number, onLaterPages=pdf_page_number)
    return output_path


def build_markdown(report_text: dict[str, list[str]]) -> Path:
    output_path = REPORT_DIR / "data607_project_report_pptx_metrics.md"
    lines: list[str] = [
        f"# {TITLE}",
        "",
        AUTHORS,
        "",
        COURSE,
        "",
        REPORT_DATE,
        "",
        "## Abstract",
        "",
        report_text["abstract"][0],
        "",
        "## Background and Introduction",
        "",
    ]
    for paragraph in report_text["introduction"]:
        lines.extend([paragraph, ""])

    lines.extend(["## Data Source and Preparation", ""])
    for paragraph in report_text["data"]:
        lines.extend([paragraph, ""])
    lines.extend(
        [
            "| Model | Features | Rows | Train | Validation | Test | Threshold |",
            "| --- | ---: | ---: | ---: | ---: | ---: | ---: |",
        ]
    )
    for row in PIPELINE_SUMMARY:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} | {row[5]} | {row[6]} |")
    lines.extend(["", "## Preliminary Analyses", ""])
    for paragraph in report_text["eda"]:
        lines.extend([paragraph, ""])
    for idx, figure in enumerate(FIGURE_SPECS[:3], start=1):
        lines.extend(
            [
                f"**Figure {idx}.** *{figure.title}*",
                "",
                f"![Figure {idx}](../figures/{figure.filename})",
                "",
                f"*Note.* {figure.note}",
                "",
            ]
        )

    lines.extend(["## Problem Statement and Working Hypotheses", ""])
    for paragraph in report_text["problem"]:
        lines.extend([paragraph, ""])

    lines.extend(["## Formal Analyses", ""])
    for paragraph in report_text["methods"]:
        lines.extend([paragraph, ""])

    lines.extend(["## Results and Interpretation", ""])
    for paragraph in report_text["results"]:
        lines.extend([paragraph, ""])
    lines.extend(
        [
            "| Model | Test F1 | Precision | Recall | ROC-AUC |",
            "| --- | ---: | ---: | ---: | ---: |",
        ]
    )
    for row in MODEL_SUMMARY:
        lines.append(f"| {row[0]} | {row[1]} | {row[2]} | {row[3]} | {row[4]} |")
    lines.extend(
        [
            "",
            f"**Figure 4.** *{FIGURE_SPECS[3].title}*",
            "",
            f"![Figure 4](../figures/{FIGURE_SPECS[3].filename})",
            "",
            f"*Note.* {FIGURE_SPECS[3].note}",
            "",
            "## Conclusions",
            "",
        ]
    )
    for paragraph in report_text["conclusion"]:
        lines.extend([paragraph, ""])

    lines.extend(["## Task Division", ""])
    for paragraph in report_text["task_division"]:
        lines.extend([paragraph, ""])

    lines.extend(["## References", ""])
    for reference in REFERENCES:
        lines.extend([f"- {reference}", ""])

    output_path.write_text("\n".join(lines), encoding="utf-8")
    return output_path


def build_report_text(df: pd.DataFrame) -> dict[str, list[str]]:
    yearly_rates = (
        df.groupby("year")["spike_lead_2"].mean().mul(100).round(2).to_dict()
    )

    abstract = (
        "Extreme hourly price spikes in Alberta's energy-only electricity market create operational and financial "
        "risk for market participants, which makes short-horizon spike prediction a relevant applied machine "
        "learning problem. This report evaluates whether publicly observable AESO market and generation data can "
        "classify whether the Alberta pool price will exceed CAD 200/MWh at t+2. Two public AESO datasets were "
        "merged at the hourly level into a file with 48,887 hourly observations covering January 2020 to July "
        "2025, while the LSTM and CNN notebook variants operated on 48,839 rows after their preprocessing steps. "
        "Exploratory analysis shows a heavy right tail in pool prices and a "
        "consistent scarcity signature: future spike hours are associated with higher current prices, higher net "
        "load, lower wind output, lower renewable share, and thinner reserve margins. Three neural-network "
        "classifiers were compared: a multilayer perceptron (MLP), a long short-term memory network (LSTM), and "
        "a one-dimensional convolutional neural network (CNN). Using the rounded metrics presented in "
        "Data607_EnergySpike_Presentation_v2.pptx, the CNN achieved the strongest thresholded test performance "
        "(F1 = 0.416, precision = 0.340, recall = 0.535, ROC-AUC = 0.941), followed by the LSTM "
        "(F1 = 0.394, precision = 0.311, recall = 0.535, ROC-AUC = 0.945) and the MLP "
        "(F1 = 0.363, precision = 0.275, recall = 0.535, ROC-AUC = 0.934). The results indicate that sequence-aware "
        "neural architectures can improve short-horizon spike classification, although the gains remain modest "
        "because the task is rare-event prediction in a market whose structure changes over time."
    )

    introduction = [
        "Alberta's wholesale electricity market is an energy-only market in which the hourly pool price is formed "
        "by the interaction of supply and demand. Because electricity cannot be stored economically at scale and "
        "system balance must be maintained continuously, periods of tight capacity, high demand, weak renewable "
        "output, or intertie stress can produce abrupt price spikes. These episodes matter for generators, large "
        "industrial consumers, retailers, and system operators because they alter dispatch incentives, hedging "
        "costs, and operating risk.",
        "The present project asks whether public system-level information is sufficient to anticipate short-run "
        "spike risk. That question is practical as well as methodological: if a model can identify high-risk "
        "hours before scarcity pricing materializes, market participants can adapt bidding, dispatch, or load "
        "scheduling decisions accordingly. Alberta is also an attractive applied setting because its market is "
        "province-specific, publicly documented, and not as overused in classroom projects as more generic "
        "benchmark datasets.",
        "The modeling choice is motivated by prior work showing that electricity prices are nonlinear, seasonal, "
        "and spike-prone. Lago et al. (2018) show that deep-learning methods often outperform traditional price "
        "forecasting approaches. Alberta-specific studies likewise motivate this comparison: Manfre Jaimes et al. "
        "(2023) use neural methods for multi-day price forecasting, while Zamudio Lopez et al. (2024) examine "
        "spike occurrence directly. In that context, the present report emphasizes transparent comparison across "
        "three neural architectures rather than claiming a production-ready forecasting system."
    ]

    data = [
        "The report uses two publicly available AESO datasets. The first is the Hourly Metered Volumes and Pool "
        "Price and AIL file, which provides pool price, Alberta Internal Load, and intertie flow information "
        "(Alberta Electric System Operator, 2025a). The second is the Historical Generation Data (CSD), which "
        "records hourly generation and system capability by fuel type (Alberta Electric System Operator, 2025b). "
        "Because both sources are distributed publicly by AESO, the group had permission to use them for academic "
        "analysis under the operator's public reporting framework.",
        "The two datasets were aligned on a common hourly Mountain Prevailing Time axis, aggregated to the "
        "province level where necessary, and then merged into a unified modeling table. The merged AESO file used "
        "for exploratory analysis contains 48,887 rows and 108 columns. The model notebooks then apply slightly "
        "different preprocessing choices: the MLP notebook reports 59 features and retains 48,887 rows, whereas "
        "the LSTM and CNN notebooks report 32 and 92 features respectively and operate on 48,839 rows.",
        "To avoid another mismatch between exploratory material and the model write-up, this report uses the "
        "presentation deck for the final rounded performance metrics and the notebooks for model-specific details "
        "such as feature counts, selected thresholds, and pipeline row counts."
    ]

    yearly_text = ", ".join(f"{year} = {rate:.2f}%" for year, rate in yearly_rates.items())
    eda = [
        "Exploratory analysis shows that Alberta pool prices are strongly right-skewed rather than approximately "
        "Gaussian. In the source data, the median hourly price is CAD 45.34/MWh, the 95th percentile is "
        "CAD 431.00/MWh, and the observed maximum is CAD 999.99/MWh. This heavy upper tail justifies a spike "
        "classification framing rather than a narrow focus on average price behavior.",
        "Temporal structure is also visible. Mean prices rise most sharply in late afternoon and early evening, "
        "and higher-price months cluster in the summer period. Year-level spike prevalence is similarly uneven: "
        f"{yearly_text}. This variability indicates that the market regime is not stationary across the full sample.",
        "The descriptive comparisons in the final deck also support a scarcity narrative. Relative to hours that "
        "are not followed by a spike, hours followed by a spike at t+2 have substantially higher current prices "
        "(CAD 366.57 versus CAD 65.02), higher net load (9,599 MW versus 8,518 MW), lower wind output "
        "(407 MW versus 1,103 MW), lower renewables share (6.7% versus 13.0%), and lower reserve margin "
        "(0.722 versus 0.851). These patterns are economically coherent because price spikes tend to emerge when "
        "dispatchable supply must absorb more load while renewable support weakens."
    ]

    problem = [
        "The central problem statement is whether publicly observable market and generation conditions contain "
        "enough information to identify a future price spike before it occurs. In operational terms, the "
        "presentation frames the task as a short-horizon warning problem; in the modeling pipeline itself, the "
        "binary target is whether the pool price exceeds CAD 200/MWh at t+2.",
        "Two working hypotheses guided the analysis. First, sequence-aware models should outperform a flat "
        "multilayer perceptron because short-run ramp events and evolving system tightness are inherently temporal. "
        "Second, future spike hours should be associated with higher load pressure and lower renewable availability, "
        "which should make variables such as current price, net load, reserve margin, and wind output informative "
        "predictors of the target."
    ]

    methods = [
        "All models were trained and evaluated with time-ordered data splits in order to avoid look-ahead bias. "
        "The final workflow used a chronological train-validation-test partition, coupled with TimeSeriesSplit "
        "cross-validation inside the pre-test horizon. Because spike hours are relatively rare, F1-score was used "
        "as the primary metric, with precision, recall, PR-AUC, and ROC-AUC used to qualify model behaviour.",
        "The analytical storyline began with a simple MLP baseline. The MLP notebook used 59 tabular features, "
        "selected a validation threshold of 0.91, and produced the weakest final F1 among the three models. This "
        "baseline was important pedagogically because it established how much could be learned from a flat feature "
        "vector before adding explicit temporal structure.",
        "The LSTM and CNN then introduced sequence modelling. The LSTM notebook used 32 features, removed manual "
        "lag variables that were redundant in a recurrent architecture, and selected a validation threshold of 0.94. "
        "The CNN notebook used 92 features with a 24-hour lookback window, selected a validation threshold of 0.77, "
        "and treated convolutional filters as detectors of local pre-spike patterns.",
        "In practical terms, the main comparison in the project is therefore not between a naive heuristic and the "
        "final CNN, but between a simple MLP baseline and temporally structured neural models, especially the CNN."
    ]

    results = [
        "Using the rounded values shown in the presentation deck, test F1 rises monotonically from the simple MLP "
        "(0.363) to the LSTM (0.394) and then to the CNN (0.416). That progression supports the core modelling "
        "claim of the project: adding temporal structure helps, and the CNN is the strongest of the three final "
        "models on the thresholded spike-classification task.",
        "The improvement is driven more by precision than by recall. The deck reports essentially the same recall "
        "for all three neural models, approximately 0.535, while precision improves from 0.275 for the MLP to "
        "0.311 for the LSTM and 0.340 for the CNN. In operational terms, the sequence-aware models do not capture "
        "more spikes than the baseline; rather, they issue fewer false alarms while preserving the same hit rate.",
        "ROC-AUC values remain high across the three models, with 0.934 for the MLP, 0.945 for the LSTM, and 0.941 "
        "for the CNN. This pattern is instructive. The LSTM slightly outperforms the CNN as a ranker of spike risk, "
        "but the CNN achieves the best thresholded F1 once the decision rule is fixed. That distinction is important "
        "because the project's applied objective is not just ranking hours by risk, but making a usable classification "
        "decision under class imbalance."
    ]

    conclusion = [
        "This project shows that publicly available AESO data contain meaningful information about short-horizon "
        "electricity price spike risk in Alberta. The formal comparison across an MLP, an LSTM, and a CNN shows "
        "that temporally structured neural models are preferable to a simple tabular baseline, with the CNN emerging "
        "as the best overall model on the final thresholded F1 metric reported in the presentation.",
        "At the same time, the results should be interpreted carefully. The sample exhibits pronounced class "
        "imbalance, the market changed materially over the 2020-2025 period, and the final gains over the MLP "
        "baseline remain limited in absolute terms. Future work should therefore extend the predictor set with weather forecasts, "
        "outage information, and offer-stack or merit-order variables, while also evaluating longer forecast "
        "horizons and cost-sensitive thresholds that better reflect operational priorities."
    ]

    task_division = [
        "Jorge Gutierrez Barajas led the LSTM workflow, including recurrent-model configuration, threshold-tuning "
        "experiments, and interpretation of the sequence-model results. Alejandro Jose Alvarado Barrera led the "
        "CNN workflow, including convolutional architecture design, class-weight calibration, and interpretation of "
        "the final best-performing model. Aravindh Palaniguru led the MLP baseline, the initial feature-engineering "
        "workflow, and the baseline comparison logic. All three members contributed jointly to data acquisition, "
        "data cleaning, exploratory analysis, interpretation of findings, and preparation of the final presentation "
        "and report."
    ]

    return {
        "abstract": [abstract],
        "introduction": introduction,
        "data": data,
        "eda": eda,
        "problem": problem,
        "methods": methods,
        "results": results,
        "conclusion": conclusion,
        "task_division": task_division,
    }


def main() -> None:
    df = load_dataframe()
    generate_figures(df)
    report_text = build_report_text(df)
    md_path = build_markdown(report_text)
    docx_path = build_docx(report_text)
    pdf_path = build_pdf(report_text)

    print(f"Markdown report: {md_path}")
    print(f"DOCX report: {docx_path}")
    print(f"PDF report: {pdf_path}")
    for figure in FIGURE_SPECS:
        print(f"Figure: {figure.path}")


if __name__ == "__main__":
    main()

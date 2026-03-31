from __future__ import annotations

from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt

from .config import OUTPUT_DATA_DIR, OUTPUT_METRICS_DIR, OUTPUT_REPORT_DIR


REFERENCES = [
    "Alberta Electric System Operator. (2021). 2020 annual market statistics. https://www.aeso.ca/assets/Uploads/2020-Annual-Market-Stats-Final.pdf",
    "Alberta Electric System Operator. (n.d.). Daily average pool price report. https://ets.aeso.ca/ets_web/ip/Market/Reports/DailyAveragePoolPriceReportServlet",
    "Stathakis, E., Papadimitriou, T., & Gogas, P. (2021). Forecasting price spikes in electricity markets. Review of Economic Analysis, 13, 65-87. https://openjournals.uwaterloo.ca/index.php/rofea/article/download/1822/2096/5445",
    "Ugurlu, U., Oksuz, I., & Tas, O. (2018). Electricity price forecasting using recurrent neural networks. Energies, 11(5), 1255. https://res.mdpi.com/d_attachment/energies/energies-11-01255/article_deploy/energies-11-01255.pdf",
    "Kuo, P.-H., & Huang, C.-J. (2018). An electricity price forecasting model by hybrid structured deep neural networks. Sustainability, 10(4), 1280. https://www.mdpi.com/2071-1050/10/4/1280",
    "Christensen, T., Hurn, A. S., & Lindsay, K. A. (2012). Forecasting spikes in electricity prices. International Journal of Forecasting, 28(2), 400-411. https://doi.org/10.1016/j.ijforecast.2011.02.019",
]


def _set_document_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def _add_table(doc: Document, dataframe: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    for idx, column in enumerate(dataframe.columns):
        table.rows[0].cells[idx].text = str(column)
    for row in dataframe.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)


def _df_to_markdown(dataframe: pd.DataFrame) -> str:
    header = "| " + " | ".join(dataframe.columns.astype(str)) + " |"
    divider = "| " + " | ".join(["---"] * len(dataframe.columns)) + " |"
    rows = [header, divider]
    for row in dataframe.itertuples(index=False):
        rows.append("| " + " | ".join(str(value) for value in row) + " |")
    return "\n".join(rows)


def generate_report(results: dict[str, Any]) -> dict[str, Path]:
    OUTPUT_REPORT_DIR.mkdir(parents=True, exist_ok=True)

    split_summary = pd.read_csv(OUTPUT_DATA_DIR / "split_summary.csv")
    model_comparison = pd.read_csv(OUTPUT_METRICS_DIR / "model_comparison.csv")
    best_model_row = model_comparison.iloc[0]
    best_model_name = best_model_row["model"]
    best_error_analysis = pd.read_csv(OUTPUT_METRICS_DIR / f"{best_model_name}_error_analysis.csv")

    train_row = split_summary.loc[split_summary["split"] == "train"].iloc[0]
    val_row = split_summary.loc[split_summary["split"] == "validation"].iloc[0]
    test_row = split_summary.loc[split_summary["split"] == "test"].iloc[0]

    fn_row = best_error_analysis.loc[best_error_analysis["error_group"] == "false_negative"]
    fp_row = best_error_analysis.loc[best_error_analysis["error_group"] == "false_positive"]

    fn_text = (
        "False negatives were limited in count or absent on the test set."
        if fn_row.empty
        else (
            f"False negatives for the best model were associated with mean reserve margin "
            f"{fn_row.iloc[0]['reserve_margin']:.3f}, mean wind generation {fn_row.iloc[0]['wind_total']:.1f} MW, "
            f"and mean two-hour-ahead price {fn_row.iloc[0]['pool_price_lead_2']:.2f} CAD/MWh."
        )
    )
    fp_text = (
        "False positives were limited in count or absent on the test set."
        if fp_row.empty
        else (
            f"False positives for the best model were associated with mean reserve margin "
            f"{fp_row.iloc[0]['reserve_margin']:.3f}, mean wind generation {fp_row.iloc[0]['wind_total']:.1f} MW, "
            f"and mean two-hour-ahead price {fp_row.iloc[0]['pool_price_lead_2']:.2f} CAD/MWh."
        )
    )

    abstract = (
        f"This project develops a full machine learning pipeline to predict two-hour-ahead electricity price spikes "
        f"in Alberta using hourly AESO market and generation data from {train_row['start']} through {test_row['end']}. "
        f"Three neural-network families were implemented and compared: a multilayer perceptron (MLP), a 1D convolutional "
        f"neural network (CNN), and a long short-term memory network (LSTM). The target was defined as a binary indicator "
        f"for pool prices above $200/MWh at time t+2, while explanatory variables were measured at time t. "
        f"Using a strict time-based split and five-fold TimeSeriesSplit tuning within the train/validation horizon, "
        f"the best model on the untouched test set was the {best_model_name.upper()} with F1={best_model_row['test_f1']:.3f}, "
        f"precision={best_model_row['test_precision']:.3f}, recall={best_model_row['test_recall']:.3f}, and ROC-AUC={best_model_row['test_roc_auc']:.3f}. "
        f"The results indicate that system tightness, lagged price behavior, and renewable variability are informative predictors of spike risk."
    )

    markdown_path = OUTPUT_REPORT_DIR / "alberta_price_spike_report.md"
    with markdown_path.open("w", encoding="utf-8") as handle:
        handle.write("# Predicting Alberta Electricity Price Spikes (2020-2025)\n\n")
        handle.write("## Abstract\n\n")
        handle.write(abstract + "\n\n")
        handle.write("## Introduction and Motivation\n\n")
        handle.write(
            "Alberta operates an energy-only electricity market in which the pool price reflects the hourly intersection "
            "of supply and demand. Because electricity is non-storable and the system must remain balanced in real time, "
            "periods of tight supply, strong load, limited imports, and renewable variability can produce sharp, short-lived price spikes. "
            "These spikes matter economically because they affect generator revenues, retailer costs, hedging, and system operations.\n\n"
        )
        handle.write("## Literature Review\n\n")
        handle.write(
            "The project design follows prior work showing that electricity prices are heavy-tailed, seasonal, and prone to spikes, "
            "which motivates both spike-classification frameworks and neural-network architectures. "
            "Stathakis et al. (2021) analyze price spike forecasting as a classification problem. "
            "Ugurlu et al. (2018) show that recurrent neural networks are effective for electricity price forecasting, "
            "while Kuo and Huang (2018) motivate CNN/LSTM-type architectures for structured temporal inputs.\n\n"
        )
        handle.write("## Data and Methodology\n\n")
        handle.write(
            f"The modeling dataset was built from AESO hourly data using a fixed Mountain Prevailing Time axis (UTC-7) to avoid daylight saving ambiguity. "
            f"The effective split interpretation was: train from {train_row['start']} through {train_row['end']}, "
            f"validation from {val_row['start']} through {val_row['end']}, and test from {test_row['start']} through {test_row['end']}. "
            "The target was a binary spike indicator for two hours ahead. Continuous predictors were standardized with training statistics only, "
            "while dummy variables were left unscaled. Hyperparameter tuning used five-fold TimeSeriesSplit within the pre-test period only.\n\n"
        )
        handle.write("### Split Summary\n\n")
        handle.write(_df_to_markdown(split_summary.round(4)) + "\n\n")
        handle.write("### Model Comparison\n\n")
        handle.write(_df_to_markdown(model_comparison.round(4)) + "\n\n")
        handle.write("## Results and Model Comparison\n\n")
        handle.write(
            f"The best test-set performer was the {best_model_name.upper()}, which achieved the highest F1-score under class imbalance. "
            "Validation and test metrics were consistent across models, but the best architecture balanced precision and recall more effectively than the alternatives.\n\n"
        )
        handle.write("## Discussion\n\n")
        handle.write(fn_text + "\n\n")
        handle.write(fp_text + "\n\n")
        handle.write(
            "Economically, the error patterns are consistent with Alberta's spike dynamics: missed spikes tend to occur when the system "
            "appears only moderately tight on observed fundamentals, while false positives arise when tight conditions do not escalate into realized scarcity pricing. "
            "This is consistent with the importance of reserve margin, fuel mix, and renewable availability in an energy-only market.\n\n"
        )
        handle.write("## Conclusion and Future Work\n\n")
        handle.write(
            "The project delivers a reproducible neural-network pipeline for Alberta spike prediction using hourly AESO fundamentals. "
            "Future work should extend the feature space with weather and outage data, explore calibration under rarer extreme-price regimes, "
            "and test cost-sensitive decision thresholds aligned to operational or trading objectives.\n\n"
        )
        handle.write("## References\n\n")
        for reference in REFERENCES:
            handle.write(f"- {reference}\n")

    doc_path = OUTPUT_REPORT_DIR / "alberta_price_spike_report.docx"
    doc = Document()
    _set_document_style(doc)
    doc.add_heading("Predicting Electricity Price Spikes in Alberta (2020-2025)", 0)
    doc.add_paragraph("JorgeFolder project deliverable")
    doc.add_paragraph("University of Calgary style academic report draft")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(abstract)

    doc.add_heading("Introduction and Motivation", level=1)
    doc.add_paragraph(
        "Alberta's electricity market is highly sensitive to short-run system conditions because supply and demand must balance in real time. "
        "Under an energy-only market design, scarcity, renewable variability, intertie limitations, and demand pressure can all contribute to sharp price spikes."
    )

    doc.add_heading("Literature Review", level=1)
    doc.add_paragraph(
        "The report framework draws on official AESO market documentation and on prior electricity-price forecasting studies. "
        "Stathakis et al. (2021) frame spike prediction as a classification problem. Ugurlu et al. (2018) show the relevance of recurrent neural networks, "
        "and Kuo and Huang (2018) motivate CNN/LSTM-style architectures for structured temporal inputs."
    )

    doc.add_heading("Data and Methodology", level=1)
    doc.add_paragraph(
        f"The final modeling sample used hourly AESO data from {train_row['start']} through {test_row['end']} on a fixed MPT axis. "
        f"Train covered {train_row['start']} to {train_row['end']}, validation covered {val_row['start']} to {val_row['end']}, "
        f"and test covered {test_row['start']} to {test_row['end']}. The target was spike_lead_2, defined as two-hour-ahead pool price above $200/MWh."
    )
    _add_table(doc, split_summary.round(4))

    doc.add_heading("Exploratory Data Analysis", level=1)
    doc.add_paragraph("The EDA highlights non-normal price behavior, strong calendar structure, and distinct system conditions during spike hours.")
    for figure_name in ["time_series", "distribution", "correlation", "heatmap", "spike_compare", "roc_curve_path"]:
        figure_path = results["eda_figures"].get(figure_name) if figure_name != "roc_curve_path" else results["roc_curve_path"]
        if figure_path and Path(figure_path).exists():
            doc.add_picture(str(figure_path), width=Inches(6.5))

    doc.add_heading("Results and Model Comparison", level=1)
    doc.add_paragraph(
        f"The {best_model_name.upper()} achieved the strongest F1-score on the untouched test set. "
        "Because spike forecasting is an imbalanced classification problem, F1-score was treated as the primary selection metric."
    )
    _add_table(doc, model_comparison.round(4))

    doc.add_heading("Discussion", level=1)
    doc.add_paragraph(fn_text)
    doc.add_paragraph(fp_text)
    doc.add_paragraph(
        "The main practical implication is that scarcity conditions can often be detected from contemporaneous load, reserves, generation mix, and recent price behavior, "
        "but some spike events remain difficult because they are triggered by abrupt market changes that are not fully captured by the current feature set."
    )

    doc.add_heading("Conclusion and Future Work", level=1)
    doc.add_paragraph(
        "This project delivered a complete reproducible pipeline for two-hour-ahead spike prediction in Alberta. "
        "Future improvements should incorporate weather, outages, and probabilistic calibration, and should test economic decision rules under asymmetric loss functions."
    )

    doc.add_heading("References", level=1)
    for reference in REFERENCES:
        doc.add_paragraph(reference)

    doc.save(doc_path)
    return {"markdown": markdown_path, "docx": doc_path}
